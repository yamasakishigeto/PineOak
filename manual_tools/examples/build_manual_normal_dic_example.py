# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

IMG = r"C:\Users\yamas\AppData\Local\Temp\claude\E--Onedrive-----------PineOak\36a5a314-6cef-4e3d-a06f-ebff22e7c6ff\scratchpad\screenshots"
OUT = r"C:\Users\yamas\AppData\Local\Temp\claude\E--Onedrive-----------PineOak\36a5a314-6cef-4e3d-a06f-ebff22e7c6ff\scratchpad\PineOak_NormalDIC_マニュアル.docx"
FONT = "游ゴシック"

doc = Document()

# ---- 既定フォント設定（日本語） ----
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), FONT)

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(22)
section.right_margin = Mm(22)


def set_east_asian(run, font=FONT):
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)


def add_title(text, size=26, color="1F4E79", align="center", space_before=6, space_after=4):
    para = doc.add_paragraph()
    para.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asian(run)
    return para


def h1(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    set_east_asian(run)
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F4E79')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return para


def h2(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2E75B6")
    set_east_asian(run)
    return para


def p(text, size=10.5, color="222222", italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asian(run)
    return para


def note(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(10)
    para.paragraph_format.left_indent = Mm(4)
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), '2E75B6')
    pbdr.append(left)
    pPr.append(pbdr)
    run = para.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")
    set_east_asian(run)
    return para


def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    run = para.runs[0] if para.runs else para.add_run("")
    run.text = text
    run.font.size = Pt(10.5)
    set_east_asian(run)
    return para


def figure(filename, target_w_mm, caption):
    path = f"{IMG}\\{filename}"
    with Image.open(path) as im:
        w, h = im.size
    target_h_mm = target_w_mm * h / w
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(path, width=Mm(target_w_mm), height=Mm(target_h_mm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    crun = cap.add_run(caption)
    crun.font.size = Pt(9)
    crun.font.color.rgb = RGBColor.from_string("666666")
    set_east_asian(crun)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def param_table(rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    widths = [Cm(3.4), Cm(2.6), Cm(9.0)]
    set_repeat_header(table.rows[0])
    hdr = table.rows[0].cells
    for i, t in enumerate(["項目", "初期値", "説明"]):
        hdr[i].width = widths[i]
        para = hdr[i].paragraphs[0]
        run = para.add_run(t)
        run.font.bold = True
        run.font.size = Pt(10)
        set_east_asian(run)
    for r in rows:
        row = table.add_row().cells
        for i, t in enumerate(r):
            row[i].width = widths[i]
            para = row[i].paragraphs[0]
            run = para.add_run(t)
            run.font.size = Pt(9.5)
            set_east_asian(run)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ============================================================
# 表紙
# ============================================================
doc.add_paragraph().paragraph_format.space_before = Pt(40)
add_title("PineOak DIC-EBSD Suite", size=28)
add_title("操作マニュアル（試作版）", size=16, color="444444", space_after=4)
add_title("Step 3 — Normal DIC（通常DIC解析）", size=13, color="2E75B6", space_after=20)
note("※ このマニュアルは画面構成の説明を目的とした試作版です。実データを使った解析結果は含まれていません。")

# ============================================================
# 1. 概要
# ============================================================
h1("1. Normal DIC とは")
p("変形前後の2枚のSEM画像を比較して、画像を細かい格子（サブセット）に分割し、"
  "各点の変位（u, v）とひずみ（εxx, εyy, εxy など）を計算する機能です。"
  "PineOakの解析ワークフローの中でも中心となるステップで、ここで得られる dic_results.xlsx が、"
  "以降のEBSDジオリファレンスや応力-ひずみマッパーの入力になります。")

# ============================================================
# 2. 起動方法
# ============================================================
h1("2. 起動方法")
p("DIC-EBSD_Suite フォルダで main.py を実行すると、ブラウザ上にランチャー画面が開きます。"
  "作業フォルダを選んだ上で、「Normal DIC」の「起動」ボタンを押すと、専用のウィザード画面が別ウィンドウで開きます。")
figure("01_launcher.png", 70, "図1: 起動直後のランチャー画面。8つのステップが縦に並んでいる")

# ============================================================
# 3. 画面構成
# ============================================================
h1("3. Normal DIC ウィザードの画面構成")
p("ウィザードは左側の「ファイル選択パネル」と、右側の「パラメータ設定パネル」の2カラム構成です。"
  "上から順に設定していけば解析まで進められるようになっています。")

h2("3.1 ① ファイル選択 ／ ② REF・DEF 選択")
p("画面左側のパネルです。まず「画像フォルダ」でSEM画像が入ったフォルダを選びます。"
  "事前にSEM Alignment（Step 2）で位置合わせをしている場合は、「Alignment JSON」でその結果ファイルも指定します（省略も可）。"
  "フォルダを選ぶと画像が一覧表示され、そこから REF（変形前・基準画像）と DEF（変形後・比較したい画像、複数選択可）を"
  "タブで切り替えて選択します。")
figure("03_dic_left_panel_crop.png", 95, "図2: 左パネル（①ファイル選択・②REF/DEF選択）")

h2("3.2 トリミング設定")
p("SEM画像の下部などに写り込む情報バー（スケールバーや撮影条件の表示帯）を解析対象から除外するための設定です。"
  "上下左右それぞれ何ピクセル切り取るかを指定し、「プレビュー表示」で実際に切り取られる範囲を画像上で確認できます。")
figure("04_dic_section_0.png", 130, "図3: トリミング設定")

h2("3.3 ③ DICパラメータ")
p("解析の中心となる設定です。PineOakのDICは「粗い探索（Stage 1）」で大まかな変位を掴んだ後、"
  "「精密探索（Stage 2）」でサブピクセル精度まで追い込む2段階方式になっています。")
figure("04_dic_section_1.png", 130, "図4: ③DICパラメータ（Stage1・Stage2・共通設定など）")
param_table([
    ["Stage1 step", "60 px", "粗い探索での格子間隔。値が小さいほど密に計算するが遅くなる"],
    ["Stage1 search", "自動(+15px)", "粗い探索の探索範囲。「グローバルシフト+○px」なら画像全体のずれを自動推定して余裕を加算、「固定値」なら常に一定範囲を探索"],
    ["Stage2 step", "15 px", "精密探索（実際に結果として出る点）の格子間隔＝解析の解像度"],
    ["Stage2 search", "5 px", "精密探索の探索範囲。Stage1の結果を初期値にするので狭くて済む"],
    ["サブセットサイズ", "31 px", "変位を求める際に比較する正方形領域の一辺。奇数を推奨"],
    ["NCCマスク閾値", "0.20", "相関値がこの値未満の点は信頼できないとしてNaN（欠損）扱いにする"],
    ["サブピクセル補間", "放物線フィット", "整数画素より細かい変位を求める方法。ガウス／2Dスプラインはより高精度だが低速"],
    ["ゲージ長さ", "1倍", "ひずみ計算に使う近傍点の間隔（1＝隣接点同士の差分）"],
    ["フレーム間隔Δt", "省略可", "入力するとひずみ速度（1秒あたりのひずみ変化）も計算する"],
    ["ひずみ種類", "微小ひずみ", "通常はこちら。ひずみが5%を超える大変形域ではグリーン-ラグランジェひずみを選ぶ"],
    ["ワーカー数", "自動(論理コア-1)", "並列計算に使うCPUスレッド数"],
])

h2("3.4 ④ 事前スキャン（search範囲の確認）")
p("本解析の前に、変形が最も大きい最後のDEF画像に対して試しに粗い探索を行い、"
  "「Stage1のsearch範囲がこの変形量に対して十分か」を確認する機能です。"
  "推奨値がStage1の設定より大きい場合は警告が出るので、その場合は「推奨値をStage1 searchに反映」ボタンで"
  "設定を直接更新できます。")
figure("04_dic_section_2.png", 130, "図5: ④事前スキャン")

h2("3.5 ⑤ スペックル品質チェック")
p("DICはSEM画像表面のランダムな模様（スペックルパターン）の濃淡を追跡して変位を求めるため、"
  "模様の質が解析精度を左右します。この機能はREF画像を解析し、模様の細かさ（スペックル半径）・"
  "コントラスト（MIG）・追跡のしやすさ（SSSIG）を数値化して、解析に十分な画質かを判定します。")
figure("04_dic_section_3.png", 130, "図6: ⑤スペックル品質チェック")

h2("3.6 ⑥ カラースケール（空欄 = 自動）")
p("解析結果のマップを表示・保存する際の、変位・ひずみ各成分のカラーバーの範囲（min/max）と配色（カラーマップ）を"
  "あらかじめ指定できます。空欄のままなら実際のデータ範囲から自動的に設定されます。"
  "以前保存した dic_config.txt を読み込んで設定を再利用することもできます。")
figure("04_dic_section_4.png", 130, "図7: ⑥カラースケール設定")

h2("3.7 実行・保存ボタン（画面下部）")
p("設定が終わったら画面下部のボタンで解析を実行します。")
figure("05_dic_footer.png", 130, "図8: 実行・キャンセル・再描画・保存ボタン")
bullet("実行 … 設定した内容でDIC解析を開始する（進捗はランチャー側のLOGに表示）")
bullet("キャンセル … 実行中の解析を中断する")
bullet("再描画 … 計算済みの結果を、カラースケールなどの表示設定だけ変えて再表示する")
bullet("ひずみ再計算・再描画 … ゲージ長さやひずみ種類を変更し、変位データからひずみだけ再計算する")
bullet("マップ・データ保存 … 結果をPNG画像とdic_results.xlsxとして保存する")

# ============================================================
# 4. まとめ
# ============================================================
h1("4. 操作の流れ（まとめ）")
bullet("① 画像フォルダを選び、必要ならAlignment JSONを指定する")
bullet("② REF（変形前）とDEF（変形後、複数可）を選ぶ")
bullet("必要ならトリミング範囲を設定してプレビューで確認する")
bullet("③ DICパラメータを設定する（迷ったら初期値のままでよい）")
bullet("④ 事前スキャンを実行し、Stage1のsearch範囲が十分か確認する")
bullet("⑤ スペックル品質チェックで画質を確認する（任意）")
bullet("「実行」ボタンで解析を開始する")
bullet("結果を確認し、必要ならカラースケールを調整して「マップ・データ保存」で確定する")

doc.save(OUT)
print("done:", OUT)
