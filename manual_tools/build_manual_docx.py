# -*- coding: utf-8 -*-
# PineOak GUIマニュアル（Word/.docx）生成テンプレート
#
# 背景: docxスキル（Node.jsのdocxパッケージ前提）を試したが、この環境には
#   Node.jsが入っておらず断念。代わりにPython+python-docxで組んでいる。
#   見た目確認は LibreOffice(soffice) も無いため、PowerShell + Word COM
#   （New-Object -ComObject Word.Application → SaveAs PDF）でPDF化し、
#   Readツールで直接PDFを読んで目視確認する（pdftoppmも無い環境）。
#
# 初回セットアップ:
#   python -m pip install python-docx pillow
#
# 使い方: このファイルをステップごとにコピーし、
#   - IMG（screenshot_capture.pyの出力フォルダ）
#   - OUT（保存先ファイル名）
#   - 本文（表紙・見出し・段落・図・表）
#   を書き換える。h1/h2/p/note/bullet/figure/param_table は使い回せる
#   共通ヘルパーなので、そのまま流用してよい。
#
# 参考: screenshot_capture.py で撮った画像をfigure()で貼り込む。
#   画像に余白（フォルダ未選択時の空欄など）が多く映る場合は、
#   PIL側で明るいピクセルの分布を見て自動トリミングしてから使うとよい
#   （Normal DICマニュアル作成時、左パネルの空欄トリミングで実施済み）。

from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

IMG = r"E:\Onedriveと同期しないフォルダ\PineOak\manual_tools\screenshots"
OUT = r"E:\Onedriveと同期しないフォルダ\PineOak\manual_tools\PineOak_マニュアル.docx"
FONT = "游ゴシック"

doc = Document()

# ---- 既定フォント設定（日本語） ----
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), FONT)

section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin = Mm(22)
section.right_margin = Mm(22)


def set_east_asian(run, font=FONT):
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)


def add_title(text, size=26, color="1F4E79", align="center", space_before=6, space_after=4):
    para = doc.add_paragraph()
    para.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asian(run)
    return para


def h1(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    set_east_asian(run)
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F4E79')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return para


def h2(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.font.size = Pt(12.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("2E75B6")
    set_east_asian(run)
    return para


def p(text, size=10.5, color="222222", italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    set_east_asian(run)
    return para


def note(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(10)
    para.paragraph_format.left_indent = Mm(4)
    pPr = para._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), '2E75B6')
    pbdr.append(left)
    pPr.append(pbdr)
    run = para.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("555555")
    set_east_asian(run)
    return para


def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    run = para.runs[0] if para.runs else para.add_run("")
    run.text = text
    run.font.size = Pt(10.5)
    set_east_asian(run)
    return para


def figure(filename, target_w_mm, caption):
    path = f"{IMG}\\{filename}"
    with Image.open(path) as im:
        w, h = im.size
    target_h_mm = target_w_mm * h / w
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run()
    run.add_picture(path, width=Mm(target_w_mm), height=Mm(target_h_mm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    crun = cap.add_run(caption)
    crun.font.size = Pt(9)
    crun.font.color.rgb = RGBColor.from_string("666666")
    set_east_asian(crun)


def set_repeat_header(row):
    """表がページをまたぐとき、見出し行を各ページ先頭に繰り返させる"""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def param_table(rows, headers=("項目", "初期値", "説明"), widths_cm=(3.4, 2.6, 9.0)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    widths = [Cm(w) for w in widths_cm]
    set_repeat_header(table.rows[0])
    hdr = table.rows[0].cells
    for i, t in enumerate(headers):
        hdr[i].width = widths[i]
        para = hdr[i].paragraphs[0]
        run = para.add_run(t)
        run.font.bold = True
        run.font.size = Pt(10)
        set_east_asian(run)
    for r in rows:
        row = table.add_row().cells
        for i, t in enumerate(r):
            row[i].width = widths[i]
            para = row[i].paragraphs[0]
            run = para.add_run(t)
            run.font.size = Pt(9.5)
            set_east_asian(run)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


# ============================================================
# ここから本文。ステップごとにコピーして書き換える。
# 以下はサンプル（Normal DICマニュアルの内容を短縮した例）。
# ============================================================
if __name__ == "__main__":
    doc.add_paragraph().paragraph_format.space_before = Pt(40)
    add_title("PineOak DIC-EBSD Suite", size=28)
    add_title("操作マニュアル（サンプル）", size=16, color="444444", space_after=4)
    note("※ これはテンプレートの動作確認用サンプルです。実際のマニュアル作成時は本文を書き換えてください。")

    h1("1. 概要")
    p("ここに、このステップで何をするかの説明を書く。")

    h1("2. 画面構成")
    h2("2.1 セクション名")
    p("ここに、screenshot_capture.pyで撮ったスクリーンショットの説明を書く。")
    # figure("section_1.png", 130, "図1: ○○セクション")
    # param_table([["項目名", "初期値", "説明"], ...])

    doc.save(OUT)
    print("done:", OUT)
